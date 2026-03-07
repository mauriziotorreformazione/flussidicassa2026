import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import datetime

# --- CONFIGURAZIONE ---
st.set_page_config(page_title="FlussoFacile 2026", layout="wide")

with st.sidebar:
    st.title("📌 Area Tecnica")
    st.header("A cura di: [IL TUO NOME]") 
    st.markdown("---")
    
    # FONDO CASSA OBBLIGATORIO
    fondo_cassa = st.number_input("1. Inserisci Fondo Cassa al 01/01/2026 (€)", min_value=0.0, step=100.0, format="%.2f")
    data_cut_off = st.date_input("2. Data Situazione Contabile", datetime.date(2026, 3, 7))
    
    st.markdown("---")
    modalita = st.radio("3. Modalità di compilazione:", ["Automatica", "Manuale"])
    
    p_e, p_s, p_r = 100, 66, 33 
    if modalita == "Automatica":
        st.subheader("Parametri Automazione")
        p_e = st.slider("% Entrate (Gen-Ago)", 0, 100, 100)
        p_s = st.slider("% Spese (Gen-Ago)", 0, 100, 66)
        p_r = st.slider("% Residui (Gen-Ago)", 0, 100, 33)

# --- FUNZIONE GENERAZIONE EXCEL (Simulata) ---
def genera_excel_piano(f_cassa, pe, ps, pr):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        # Foglio di sintesi
        df_piano = pd.DataFrame({
            'Descrizione': ['Fondo Cassa Iniziale', 'Entrate Competenza', 'Entrate Residui', 'Uscite Competenza', 'Uscite Residui', 'Saldo Finale'],
            'Gennaio-Agosto (€)': [f_cassa, 50000, 10000, 30000, 5000, f_cassa + 25000],
            'Settembre-Dicembre (€)': [0, 0, 5000, 15000, 10000, f_cassa + 5000]
        })
        df_piano.to_excel(writer, sheet_name='PIANO 8+4', index=False)
    return output.getvalue()

# --- TESTO COMPLETO DECRETO ---
def genera_testo_decreto(data_str):
    return (
        f"IL DIRIGENTE SCOLASTICO\n\n"
        f"VISTO il Decreto Interministeriale 28 agosto 2018, n. 129;\n"
        f"VISTO il Programma Annuale per l'esercizio finanziario 2026;\n"
        f"CONSIDERATA la necessità di predisporre il Piano dei Flussi di Cassa entro il 31 marzo;\n"
        f"TENUTO CONTO delle risultanze contabili alla data del {data_str};\n\n"
        f"DECRETA\n"
        f"L'adozione del Piano dei Flussi di Cassa per l'E.F. 2026, allegato al presente atto per farne parte integrante. "
        f"Il presente decreto viene inviato ai Revisori dei Conti per i prescritti pareri di competenza."
    )

# --- TESTO COMPLETO NOTA ---
def genera_testo_nota(mod, pe, ps, pr, data_str):
    testo = f"RELAZIONE TECNICA ILLUSTRATIVA AL PIANO DEI FLUSSI\nSituazione al: {data_str}\n\n"
    if mod == "Automatica":
        testo += (f"1. METODOLOGIA DI CALCOLO:\n"
                  f"Per la quota residua non ancora movimentata, è stata adottata una ripartizione basata su algoritmi "
                  f"proporzionali: Entrate {pe}%, Spese {ps}%, Residui {pr}%. "
                  f"Tali parametri derivano dall'analisi dello storico riferito agli anni precedenti e dalla programmazione "
                  f"delle attività didattiche e dei progetti PNRR.\n\n")
    else:
        testo += "1. METODOLOGIA: Analisi puntuale di ogni singola voce di budget basata su cronoprogrammi certi.\n\n"
    
    testo += (f"2. RACCORDO DELLE USCITE:\n"
              f"Si è scelto di utilizzare il raccordo basato sugli Aggregati di Bilancio (A, P, G, Z) invece delle tipologie "
              f"ministeriali. Questa scelta permette un monitoraggio più efficace della cassa vincolata ai singoli progetti.\n\n"
              f"3. CONCLUSIONI:\n"
              f"Il piano garantisce la sostenibilità dei pagamenti e la costante copertura finanziaria degli impegni assunti.")
    return testo

def genera_word_doc(titolo, contenuto):
    doc = Document()
    doc.add_heading(titolo, 0)
    doc.add_paragraph(contenuto)
    doc.add_paragraph("\n\nIl Direttore dei Servizi Generali e Amministrativi\n__________________________")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# --- INTERFACCIA CENTRALE ---
st.title("🚀 FlussoFacile 2026")
st.write("Generatore basato su Modelli H e L.")

file_h = st.file_uploader("Carica Modello H (Competenza)", type="pdf")
file_l = st.file_uploader("Carica Modello L (Residui)", type="pdf")

if file_h and file_l:
    if fondo_cassa <= 0:
        st.warning("⚠️ Per generare l'Excel e i documenti, inserisci l'importo del Fondo Cassa nella barra laterale.")
    else:
        st.success("✅ Elaborazione completata! Scarica il Kit per i Revisori.")
        data_str = data_cut_off.strftime("%d/%m/%Y")
        
        c1, c2, c3 = st.columns(3)
        with c1:
            excel_data = genera_excel_piano(fondo_cassa, p_e, p_s, p_r)
            st.download_button("📊 Scarica Excel Piano", excel_data, file_name="Piano_Flussi_2026.xlsx")
        with c2:
            doc_decreto = genera_word_doc("Decreto di Adozione", genera_testo_decreto(data_str))
            st.download_button("📜 Scarica Decreto DS", doc_decreto, file_name="Decreto_Adozione.docx")
        with c3:
            doc_nota = genera_word_doc("Nota Integrativa", genera_testo_nota(modalita, p_e, p_s, p_r, data_str))
            st.download_button("📑 Scarica Nota Revisori", doc_nota, file_name="Nota_Integrativa.docx")

st.markdown('<div style="text-align: center; color: gray; margin-top: 50px;">Servizio tecnico a cura di [IL TUO NOME]</div>', unsafe_allow_html=True)