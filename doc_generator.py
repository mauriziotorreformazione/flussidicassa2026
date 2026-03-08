"""
Generazione documenti Word:
1. Nota di accompagnamento per i Revisori dei Conti
2. Decreto di adozione del Dirigente Scolastico
"""
import io


def _build_nota_automatica(doc, dati_scuola, percentuali):
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    anno = dati_scuola.get("anno_esercizio", "2026")
    nome = dati_scuola.get("nome_istituto", "ISTITUTO")
    dirigente = dati_scuola.get("dirigente", "Il Dirigente Scolastico")
    dsga = dati_scuola.get("dsga", "Il DSGA")
    citta = dati_scuola.get("citta", "")
    data_decreto = dati_scuola.get("data_decreto", "")
    prot = dati_scuola.get("num_protocollo", "")

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nome.upper())
    run.bold = True
    run.font.size = Pt(13)

    if dati_scuola.get("indirizzo"):
        p2 = doc.add_paragraph(dati_scuola.get("indirizzo", "") + " — " + dati_scuola.get("citta", ""))
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Oggetto
    obj_p = doc.add_paragraph()
    obj_p.add_run("OGGETTO: ").bold = True
    obj_p.add_run(f"Piano Annuale dei Flussi di Cassa — Esercizio finanziario {anno} — "
                  "Nota metodologica.")

    doc.add_paragraph()

    # Collegio dei revisori
    dest_p = doc.add_paragraph()
    dest_p.add_run("Al Collegio dei Revisori dei Conti").bold = True
    dest_p.add_run("\nAll'Albo dell'Istituto")
    doc.add_paragraph()

    # Body
    doc.add_paragraph(
        f"In ottemperanza a quanto previsto dall'art. 6 del D.L. 22 ottobre 2024, n. 155, "
        f"convertito con modificazioni dalla Legge 19 dicembre 2024, n. 189, si trasmette "
        f"il Piano Annuale dei Flussi di Cassa relativo all'esercizio finanziario {anno}."
    )

    doc.add_paragraph(
        "Il documento è stato elaborato a partire dai seguenti atti contabili:"
    )
    ul = doc.add_paragraph(style="List Bullet")
    ul.add_run("Modello H+ — Conto Consuntivo (Conto Finanziario), relativo all'esercizio in corso;")
    ul2 = doc.add_paragraph(style="List Bullet")
    ul2.add_run("Modello L — Elenco Residui Attivi e Passivi.")
    doc.add_paragraph()

    # Methodology
    h = doc.add_paragraph()
    h.add_run("METODOLOGIA ADOTTATA (Modalità Automatica)").bold = True

    doc.add_paragraph(
        "La ripartizione degli importi tra il periodo gennaio–agosto e il periodo settembre–dicembre "
        "è stata effettuata applicando le percentuali di seguito indicate, concordate con le esigenze "
        "operative dell'Istituto:"
    )

    # Percentuale table
    from docx.shared import Inches
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"

    hdr = table.rows[0]
    hdr.cells[0].text = "Categoria"
    hdr.cells[1].text = "Gen–Ago"
    hdr.cells[2].text = "Set–Dic"
    for cell in hdr.cells:
        for run in cell.paragraphs[0].runs:
            run.bold = True

    rows_data = [
        ("Entrate in competenza",
         f"{percentuali.get('entrate_competenza_gen_ago', 100)}%",
         f"{percentuali.get('entrate_competenza_set_dic', 0)}%"),
        ("Spese in competenza",
         f"{percentuali.get('spese_competenza_gen_ago', 67)}%",
         f"{percentuali.get('spese_competenza_set_dic', 33)}%"),
        ("Residui attivi",
         f"{percentuali.get('residui_attivi_gen_ago', 33)}%",
         f"{percentuali.get('residui_attivi_set_dic', 67)}%"),
        ("Residui passivi",
         f"{percentuali.get('residui_passivi_gen_ago', 33)}%",
         f"{percentuali.get('residui_passivi_set_dic', 67)}%"),
    ]
    for i, (cat, ga, sd) in enumerate(rows_data, 1):
        table.rows[i].cells[0].text = cat
        table.rows[i].cells[1].text = ga
        table.rows[i].cells[2].text = sd

    doc.add_paragraph()
    doc.add_paragraph(
        "Le somme già riscosse/pagate alla data di elaborazione sono state integralmente allocate "
        "nel periodo gennaio–agosto. Le percentuali sono state applicate alla differenza tra la "
        "Programmazione Definitiva e gli importi già movimentati."
    )
    doc.add_paragraph(
        "Le voci con Programmazione Definitiva pari a zero ma con movimenti già registrati sono "
        "state evidenziate nel file Excel e interamente attribuite al periodo gennaio–agosto."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        f"{citta}, {data_decreto}" if citta else data_decreto or "_____________"
    )
    doc.add_paragraph()

    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.rows[0].cells[0].text = "Il Direttore S.G.A."
    sign_table.rows[0].cells[1].text = "Il Dirigente Scolastico"
    sign_table.rows[1].cells[0].text = dsga
    sign_table.rows[1].cells[1].text = dirigente
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _build_nota_manuale(doc, dati_scuola):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    anno = dati_scuola.get("anno_esercizio", "2026")
    nome = dati_scuola.get("nome_istituto", "ISTITUTO")
    dirigente = dati_scuola.get("dirigente", "Il Dirigente Scolastico")
    dsga = dati_scuola.get("dsga", "Il DSGA")
    citta = dati_scuola.get("citta", "")
    data_decreto = dati_scuola.get("data_decreto", "")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(nome.upper())
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    obj_p = doc.add_paragraph()
    obj_p.add_run("OGGETTO: ").bold = True
    obj_p.add_run(f"Piano Annuale dei Flussi di Cassa — Esercizio finanziario {anno} — "
                  "Nota metodologica.")

    doc.add_paragraph()
    dest_p = doc.add_paragraph()
    dest_p.add_run("Al Collegio dei Revisori dei Conti").bold = True
    doc.add_paragraph()

    doc.add_paragraph(
        f"In ottemperanza a quanto previsto dall'art. 6 del D.L. 22 ottobre 2024, n. 155, "
        f"convertito con modificazioni dalla Legge 19 dicembre 2024, n. 189, si trasmette "
        f"il Piano Annuale dei Flussi di Cassa relativo all'esercizio finanziario {anno}."
    )

    h = doc.add_paragraph()
    h.add_run("METODOLOGIA ADOTTATA (Modalità Manuale)").bold = True

    doc.add_paragraph(
        "La ripartizione degli importi tra il periodo gennaio–agosto e il periodo settembre–dicembre "
        "è stata effettuata sulla base della conoscenza diretta dei flussi di cassa attesi, "
        "tenendo conto della natura e delle scadenze delle singole voci di entrata e di spesa."
    )
    doc.add_paragraph(
        "Le somme già riscosse/pagate alla data di adozione del Piano sono state allocate "
        "nel periodo gennaio–agosto. Per le restanti voci, la ripartizione è stata determinata "
        "in base a specifici elementi di conoscenza (contratti, comunicazioni ministeriali, "
        "scadenze previste, prassi consolidata)."
    )

    doc.add_paragraph()
    doc.add_paragraph(citta + ", " + (data_decreto or "_____________") if citta else "_____________")
    doc.add_paragraph()

    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.rows[0].cells[0].text = "Il Direttore S.G.A."
    sign_table.rows[0].cells[1].text = "Il Dirigente Scolastico"
    sign_table.rows[1].cells[0].text = dsga
    sign_table.rows[1].cells[1].text = dirigente
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def genera_nota(dati_scuola, percentuali, modalita) -> bytes:
    from docx import Document
    from docx.shared import Cm

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    if modalita == "automatica":
        _build_nota_automatica(doc, dati_scuola, percentuali)
    else:
        _build_nota_manuale(doc, dati_scuola)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def genera_decreto(dati_scuola, percentuali, modalita) -> bytes:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    anno = dati_scuola.get("anno_esercizio", "2026")
    nome = dati_scuola.get("nome_istituto", "ISTITUTO")
    cf = dati_scuola.get("cf", "")
    codice_mecc = dati_scuola.get("codice_mecc", "")
    dirigente = dati_scuola.get("dirigente", "Il Dirigente Scolastico")
    dsga = dati_scuola.get("dsga", "Il DSGA")
    citta = dati_scuola.get("citta", "")
    data_decreto = dati_scuola.get("data_decreto", "")
    num_prot = dati_scuola.get("num_protocollo", "")
    data_delibera = dati_scuola.get("data_delibera_ci", "")
    num_delibera = dati_scuola.get("num_delibera_ci", "")
    fondo = dati_scuola.get("fondo_cassa", 0)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(nome.upper())
    r.bold = True
    r.font.size = Pt(13)

    if dati_scuola.get("indirizzo"):
        p2 = doc.add_paragraph(dati_scuola.get("indirizzo", "") + " — " + citta)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cf:
            p3 = doc.add_paragraph(f"C.F. {cf} — Cod. Mecc. {codice_mecc}")
            p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Protocollo
    if num_prot or data_decreto:
        p_prot = doc.add_paragraph()
        p_prot.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_prot.add_run(f"Prot. n. {num_prot} del {data_decreto}")

    doc.add_paragraph()

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run(f"DECRETO DEL DIRIGENTE SCOLASTICO\nN. {num_prot}/{anno}")
    t_run.bold = True
    t_run.font.size = Pt(12)

    obj_p = doc.add_paragraph()
    obj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    obj_p.add_run(f"OGGETTO: Adozione del Piano Annuale dei Flussi di Cassa — Esercizio {anno}.")

    doc.add_paragraph()

    # IL DIRIGENTE SCOLASTICO
    ds_header = doc.add_paragraph()
    ds_header.add_run("IL DIRIGENTE SCOLASTICO").bold = True

    # VISTI
    doc.add_paragraph(
        f"VISTO il D.I. n. 129 del 28/08/2018 recante il Regolamento concernente le istruzioni "
        f"generali sulla gestione amministrativo-contabile delle istituzioni scolastiche;"
    )
    doc.add_paragraph(
        "VISTO l'art. 6 del D.L. 22 ottobre 2024, n. 155, convertito con modificazioni dalla "
        "Legge 19 dicembre 2024, n. 189, che introduce l'obbligo di adottare il Piano Annuale "
        "dei Flussi di Cassa entro il 28 febbraio di ciascun anno;"
    )
    doc.add_paragraph(
        f"VISTO il Programma Annuale dell'esercizio {anno} approvato dal Consiglio d'Istituto;"
    )
    if num_delibera and data_delibera:
        doc.add_paragraph(
            f"VISTA la delibera n. {num_delibera} del {data_delibera} del Consiglio d'Istituto "
            f"con la quale si prende atto del Piano Annuale dei Flussi di Cassa;"
        )
    doc.add_paragraph(
        f"VISTO il Conto Finanziario (Modello H+) e l'Elenco Residui (Modello L) quale base "
        f"per la redazione del Piano;"
    )
    doc.add_paragraph(
        f"CONSIDERATO che il fondo cassa al 01/01/{anno} ammonta a € {fondo:,.2f};"
    )
    doc.add_paragraph(
        "SENTITO il Direttore dei Servizi Generali e Amministrativi;"
    )

    doc.add_paragraph()

    # DECRETA
    decreta_p = doc.add_paragraph()
    decreta_p.add_run("DECRETA").bold = True

    doc.add_paragraph("Art. 1 — È adottato il Piano Annuale dei Flussi di Cassa relativo "
                      f"all'esercizio finanziario {anno}, allegato al presente decreto quale parte integrante.")
    doc.add_paragraph("Art. 2 — Il Piano è predisposto sulla base del Modello H+ (Conto Finanziario) "
                      "e del Modello L (Elenco Residui) ed esprime la previsione dei flussi di cassa "
                      "in entrata e in uscita per i periodi gennaio–agosto e settembre–dicembre.")
    doc.add_paragraph("Art. 3 — Il presente decreto viene trasmesso al Collegio dei Revisori dei Conti "
                      "unitamente alla nota metodologica di accompagnamento, e pubblicato all'Albo "
                      "dell'Istituto.")

    doc.add_paragraph()
    doc.add_paragraph(f"{citta}, {data_decreto}" if citta and data_decreto else "_____________")
    doc.add_paragraph()

    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.rows[0].cells[0].text = "Il Direttore S.G.A."
    sign_table.rows[0].cells[1].text = "Il Dirigente Scolastico"
    sign_table.rows[1].cells[0].text = dsga
    sign_table.rows[1].cells[1].text = dirigente
    for row in sign_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
